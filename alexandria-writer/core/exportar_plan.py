#!/usr/bin/env python3
"""
Exportar Plan de Edición — Alexandria Writer v3
=================================================
Lee las decisiones del autor (hallazgos marcados como aceptado/rechazado/modificado)
y genera un documento Word con el plan de edición personalizado.

Dependencia: pip install python-docx

Uso:
  python core/exportar_plan.py --autor <AUTOR> --libro <LIBRO>
  python core/exportar_plan.py --autor <AUTOR> --libro <LIBRO> --iteracion 08
"""

import sys
import json
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional

CORE_DIR = Path(__file__).parent
sys.path.insert(0, str(CORE_DIR))

from config_v3 import ProyectoPaths


def _verificar_docx() -> bool:
    try:
        import docx
        return True
    except ImportError:
        print("   [ERROR] python-docx no instalado. Ejecutá: pip install python-docx")
        return False


def cargar_hallazgos_con_decisiones(paths: ProyectoPaths) -> List[Dict]:
    """Carga hallazgos y adjunta las decisiones del autor."""
    h_path = paths.hallazgos_path()
    d_path = paths.m1_diagnostico / "decisiones_autor.json"

    if not h_path.exists():
        print("   [ADVERTENCIA] hallazgos.json no encontrado.")
        return []

    with open(h_path, encoding="utf-8") as f:
        hallazgos = json.load(f)

    decisiones: Dict = {}
    if d_path.exists():
        with open(d_path, encoding="utf-8") as f:
            decisiones = json.load(f)

    for h in hallazgos:
        dec = decisiones.get(h.get("id", ""), {})
        h["_decision"] = dec.get("decision", "pendiente")
        h["_nota_autor"] = dec.get("nota", "")

    return hallazgos


def cargar_backlog(paths: ProyectoPaths) -> List[Dict]:
    b_path = paths.backlog_path()
    if not b_path.exists():
        return []
    with open(b_path, encoding="utf-8") as f:
        return json.load(f).get("backlog", [])


def cargar_entregable(ruta: Path) -> str:
    if ruta.exists():
        return ruta.read_text(encoding="utf-8")
    return ""


def exportar_word(
    paths: ProyectoPaths,
    hallazgos: List[Dict],
    backlog: List[Dict],
    ruta_salida: Path,
) -> bool:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ─── Estilos base ───
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(11)

    # ─── Portada ───
    titulo = doc.add_heading(f"Plan de Edición — {paths.libro_id}", 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(f"Autor: {paths.autor}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2 = doc.add_paragraph(f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M')} · Alexandria Writer v3")
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ─── Resumen ejecutivo (de M7 si existe) ───
    brief = cargar_entregable(paths.m7_output_profesional / "brief_final_ejecutivo.md")
    if brief:
        doc.add_heading("Resumen Ejecutivo", 1)
        for linea in brief.splitlines():
            if linea.startswith("# "):
                continue
            if linea.startswith("## "):
                doc.add_heading(linea[3:], 2)
            elif linea.startswith("> "):
                p = doc.add_paragraph(linea[2:])
                p.style = "Quote" if "Quote" in [s.name for s in doc.styles] else "Normal"
            elif linea.strip():
                doc.add_paragraph(linea)
        doc.add_page_break()

    # ─── Estadísticas de decisiones ───
    doc.add_heading("Estado de Decisiones", 1)
    total = len(hallazgos)
    aceptados = [h for h in hallazgos if h.get("_decision") == "aceptado"]
    rechazados = [h for h in hallazgos if h.get("_decision") == "rechazado"]
    modificados = [h for h in hallazgos if h.get("_decision") == "modificado"]
    pendientes = [h for h in hallazgos if h.get("_decision") == "pendiente"]

    tabla_stats = doc.add_table(rows=5, cols=2)
    tabla_stats.style = "Table Grid"
    celdas = [
        ("Total hallazgos", str(total)),
        ("Aceptados ✓", str(len(aceptados))),
        ("Rechazados ✗", str(len(rechazados))),
        ("Modificados ~", str(len(modificados))),
        ("Pendientes de decisión", str(len(pendientes))),
    ]
    for i, (k, v) in enumerate(celdas):
        tabla_stats.cell(i, 0).text = k
        tabla_stats.cell(i, 1).text = v
    doc.add_paragraph()

    # ─── Plan de intervención (sección principal) ───
    doc.add_heading("Plan de Intervención", 1)

    # Aceptados ordenados por rank del backlog
    backlog_rank: Dict[str, int] = {b.get("id", ""): b.get("rank", 999) for b in backlog}
    aceptados_ordenados = sorted(aceptados, key=lambda h: backlog_rank.get(h.get("id", ""), 999))

    if aceptados_ordenados:
        doc.add_heading("Cambios a implementar", 2)
        for i, h in enumerate(aceptados_ordenados, 1):
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(f"[{h.get('id')}] ")
            run.bold = True
            p.add_run(f"{h.get('tipo', '?').upper()} — Sev. {h.get('severidad', '?')}/5")

            if h.get("cita_textual"):
                cita = doc.add_paragraph(f"« {h['cita_textual'][:200]} »")
                cita.paragraph_format.left_indent = Inches(0.5)
                cita.runs[0].italic = True

            if h.get("descripcion"):
                doc.add_paragraph(f"Hallazgo: {h['descripcion'][:300]}")

            if h.get("intervencion_sugerida"):
                accion = doc.add_paragraph(f"Acción: {h['intervencion_sugerida'][:300]}")
                accion.runs[0].bold = True

            if h.get("_nota_autor"):
                doc.add_paragraph(f"Nota del autor: {h['_nota_autor']}")

            doc.add_paragraph()
    else:
        doc.add_paragraph("No hay hallazgos aceptados aún.")

    # Modificados
    if modificados:
        doc.add_heading("Cambios con modificación del autor", 2)
        for h in modificados:
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(f"[{h.get('id')}] ")
            run.bold = True
            p.add_run(f"{h.get('tipo', '?').upper()} — Sev. {h.get('severidad', '?')}/5")
            if h.get("descripcion"):
                doc.add_paragraph(f"Hallazgo original: {h['descripcion'][:200]}")
            if h.get("_nota_autor"):
                nota = doc.add_paragraph(f"Modificación del autor: {h['_nota_autor']}")
                nota.runs[0].bold = True
            doc.add_paragraph()

    doc.add_page_break()

    # ─── Lo que NO se toca (rechazados) ───
    if rechazados:
        doc.add_heading("Lo que NO se modifica en esta ronda", 1)
        doc.add_paragraph(
            "El autor rechazó las siguientes sugerencias. "
            "No deben aplicarse en esta iteración."
        )
        for h in rechazados:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"[{h.get('id')}] ")
            run.bold = True
            p.add_run(h.get("descripcion", "")[:150])
            if h.get("_nota_autor"):
                doc.add_paragraph(f"  → Razón del autor: {h['_nota_autor']}", style="List Bullet")

        doc.add_page_break()

    # ─── Diagnóstico de desarrollo (de M7 si existe) ───
    diagnostico = cargar_entregable(paths.m7_output_profesional / "diagnostico_desarrollo.md")
    if diagnostico:
        doc.add_heading("Diagnóstico Editorial", 1)
        for linea in diagnostico.splitlines():
            if linea.startswith("# "):
                continue
            elif linea.startswith("## "):
                doc.add_heading(linea[3:], 2)
            elif linea.startswith("> "):
                p = doc.add_paragraph(linea[2:])
                p.paragraph_format.left_indent = Inches(0.3)
            elif linea.strip():
                doc.add_paragraph(linea)

    # ─── Guardar ───
    ruta_salida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(ruta_salida))
    return True


def exportar_markdown(
    paths: ProyectoPaths,
    hallazgos: List[Dict],
    backlog: List[Dict],
    ruta_salida: Path,
) -> bool:
    """Versión Markdown del plan (no requiere python-docx)."""
    backlog_rank: Dict[str, int] = {b.get("id", ""): b.get("rank", 999) for b in backlog}
    aceptados = sorted(
        [h for h in hallazgos if h.get("_decision") == "aceptado"],
        key=lambda h: backlog_rank.get(h.get("id", ""), 999),
    )
    rechazados = [h for h in hallazgos if h.get("_decision") == "rechazado"]
    modificados = [h for h in hallazgos if h.get("_decision") == "modificado"]
    pendientes = [h for h in hallazgos if h.get("_decision") == "pendiente"]

    lineas = [
        f"# Plan de Edición — {paths.libro_id}\n",
        f"*Autor: {paths.autor} · Generado: {datetime.now().strftime('%Y-%m-%d %H:%M')} · Alexandria Writer v3*\n",
        "---\n",
        f"**Total hallazgos:** {len(hallazgos)} · "
        f"**Aceptados:** {len(aceptados)} · "
        f"**Rechazados:** {len(rechazados)} · "
        f"**Modificados:** {len(modificados)} · "
        f"**Pendientes:** {len(pendientes)}\n",
        "---\n",
    ]

    if aceptados:
        lineas.append("## Cambios a implementar\n")
        for i, h in enumerate(aceptados, 1):
            lineas.append(f"\n### {i}. [{h.get('id')}] {h.get('tipo', '').upper()} — Sev. {h.get('severidad', '?')}/5")
            if h.get("cita_textual"):
                lineas.append(f"\n> *{h['cita_textual'][:300]}*")
            lineas.append(f"\n**Hallazgo:** {h.get('descripcion', '')[:300]}")
            if h.get("intervencion_sugerida"):
                lineas.append(f"\n**Acción:** {h.get('intervencion_sugerida', '')[:300]}")
            if h.get("_nota_autor"):
                lineas.append(f"\n**Nota del autor:** {h['_nota_autor']}")

    if modificados:
        lineas.append("\n---\n\n## Cambios con modificación del autor\n")
        for h in modificados:
            lineas.append(f"\n### [{h.get('id')}] {h.get('descripcion', '')[:150]}")
            if h.get("_nota_autor"):
                lineas.append(f"\n**Modificación:** {h['_nota_autor']}")

    if rechazados:
        lineas.append("\n---\n\n## Lo que NO se modifica en esta ronda\n")
        for h in rechazados:
            nota = f" · Razón: {h['_nota_autor']}" if h.get("_nota_autor") else ""
            lineas.append(f"- **[{h.get('id')}]** {h.get('descripcion', '')[:150]}{nota}")

    ruta_salida.write_text("\n".join(lineas), encoding="utf-8")
    return True


def ejecutar(autor: str, libro: str, iteracion: Optional[str] = None, formato: str = "ambos") -> bool:
    paths = ProyectoPaths(autor=autor, libro_id=libro)

    print(f"\n[Exportar] {autor} / {libro}")
    hallazgos = cargar_hallazgos_con_decisiones(paths)
    backlog = cargar_backlog(paths)

    if not hallazgos:
        print("   Sin hallazgos para exportar.")
        return False

    aceptados = sum(1 for h in hallazgos if h.get("_decision") == "aceptado")
    total = len(hallazgos)
    print(f"   {total} hallazgos · {aceptados} aceptados por el autor")

    fecha = datetime.now().strftime("%Y%m%d")
    base = paths.m7_output_profesional if paths.m7_output_profesional.exists() else paths.proyecto_dir
    base.mkdir(parents=True, exist_ok=True)

    # Markdown (siempre)
    ruta_md = base / f"plan_edicion_{fecha}.md"
    if exportar_markdown(paths, hallazgos, backlog, ruta_md):
        print(f"   ✓ Markdown: {ruta_md}")

    # Word (si tiene python-docx)
    if formato in ("word", "ambos"):
        if _verificar_docx():
            ruta_docx = base / f"plan_edicion_{fecha}.docx"
            if exportar_word(paths, hallazgos, backlog, ruta_docx):
                print(f"   ✓ Word: {ruta_docx}")
            else:
                print("   [ERROR] No se pudo generar el Word.")

    return True


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Exportar plan de edición (Markdown + Word)")
    parser.add_argument("--autor", required=True)
    parser.add_argument("--libro", required=True)
    parser.add_argument("--formato", default="ambos", choices=["markdown", "word", "ambos"])
    args = parser.parse_args()
    exit(0 if ejecutar(args.autor, args.libro, formato=args.formato) else 1)
