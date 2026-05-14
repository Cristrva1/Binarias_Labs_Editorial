"""
CUESTIONARIO GENERICO DE EXPLORACION PARA ESCRITORES
Plantilla reutilizable para cualquier libro (ficcion, no ficcion, poesia, etc.)
Diseño: Dorado/Ocre + Azul Marino | Enfocado en valor editorial, no en marketing.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def add_horizontal_line(doc, color_hex="D4A843"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("\n" + "━" * 60 + "\n")
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_section_header(doc, number, title, color_hex="1B3A5C"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)

    run = p.add_run(f"SECCION {number}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16))
    run.font.name = "Calibri"

    run2 = p.add_run(f"  —  {title.upper()}")
    run2.bold = True
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16))
    run2.font.name = "Calibri"

def add_question(doc, number, question, contexto, ejemplo=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)

    run_num = p.add_run(f"{number}.  ")
    run_num.bold = True
    run_num.font.size = Pt(12)
    run_num.font.color.rgb = RGBColor(212, 168, 67)
    run_num.font.name = "Calibri"

    run_q = p.add_run(question)
    run_q.bold = True
    run_q.font.size = Pt(12)
    run_q.font.color.rgb = RGBColor(27, 58, 92)
    run_q.font.name = "Calibri"

    p_ctx = doc.add_paragraph()
    p_ctx.paragraph_format.space_before = Pt(2)
    p_ctx.paragraph_format.space_after = Pt(6)
    p_ctx.paragraph_format.left_indent = Inches(0.25)

    run_ctx = p_ctx.add_run(f"💡  {contexto}")
    run_ctx.italic = True
    run_ctx.font.size = Pt(10)
    run_ctx.font.color.rgb = RGBColor(100, 100, 100)
    run_ctx.font.name = "Calibri"

    if ejemplo:
        p_ej = doc.add_paragraph()
        p_ej.paragraph_format.space_before = Pt(2)
        p_ej.paragraph_format.space_after = Pt(4)
        p_ej.paragraph_format.left_indent = Inches(0.25)

        run_ej = p_ej.add_run(f"Ejemplo: {ejemplo}")
        run_ej.italic = True
        run_ej.font.size = Pt(10)
        run_ej.font.color.rgb = RGBColor(120, 100, 60)
        run_ej.font.name = "Calibri"

    p_resp = doc.add_paragraph()
    p_resp.paragraph_format.space_before = Pt(6)
    p_resp.paragraph_format.space_after = Pt(4)

    run_label = p_resp.add_run("Tu respuesta:")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_label.font.color.rgb = RGBColor(150, 150, 150)
    run_label.font.name = "Calibri"

    for _ in range(3):
        p_line = doc.add_paragraph()
        p_line.paragraph_format.space_before = Pt(2)
        p_line.paragraph_format.space_after = Pt(2)
        p_line.paragraph_format.left_indent = Inches(0.25)

        run_line = p_line.add_run("_" * 90)
        run_line.font.size = Pt(10)
        run_line.font.color.rgb = RGBColor(200, 200, 200)
        run_line.font.name = "Calibri"

def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run("✦  ✦  ✦")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(212, 168, 67)
    run.font.name = "Calibri"

def main():
    doc = Document()
    sections = doc.sections[0]
    sections.top_margin = Cm(2.5)
    sections.bottom_margin = Cm(2.5)
    sections.left_margin = Cm(2.5)
    sections.right_margin = Cm(2.5)

    # === PORTADA ===
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)

    run = p.add_run("CUESTIONARIO DE EXPLORACION")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(27, 58, 92)
    run.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(8)

    run2 = p2.add_run("para Escritores")
    run2.italic = True
    run2.font.size = Pt(18)
    run2.font.color.rgb = RGBColor(212, 168, 67)
    run2.font.name = "Calibri"

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(20)

    run_sub = p_sub.add_run("Ficcion  ·  No Ficcion  ·  Poesia  ·  Cualquier genero")
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = RGBColor(120, 120, 120)
    run_sub.font.name = "Calibri"

    # Caja informativa
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F0E6")

    p_info = cell.paragraphs[0]
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_info.paragraph_format.space_before = Pt(8)
    p_info.paragraph_format.space_after = Pt(8)

    run_info = p_info.add_run(
        "📖  Proposito: Ayudarte a entender tu propio manuscrito desde fuera.\n"
        "Tus respuestas alimentaran el analisis editorial automatizado,\n"
        "para que las recomendaciones del sistema resuenen con tu vision.\n\n"
        "⏱  Tiempo estimado: 25-35 minutos\n"
        "✍  Responde con la extension que necesites.\n"
        "💡 No hay respuestas incorrectas."
    )
    run_info.font.size = Pt(11)
    run_info.font.color.rgb = RGBColor(60, 60, 60)
    run_info.font.name = "Calibri"

    doc.add_paragraph()

    # === DATOS DEL LIBRO ===
    p_datos = doc.add_paragraph()
    p_datos.paragraph_format.space_before = Pt(10)
    p_datos.paragraph_format.space_after = Pt(6)

    run_datos = p_datos.add_run("TITULO DEL LIBRO:  _________________________________________________")
    run_datos.bold = True
    run_datos.font.size = Pt(12)
    run_datos.font.color.rgb = RGBColor(27, 58, 92)

    p_autor = doc.add_paragraph()
    p_autor.paragraph_format.space_after = Pt(6)

    run_autor = p_autor.add_run("NOMBRE DEL AUTOR:  _________________________________________________")
    run_autor.bold = True
    run_autor.font.size = Pt(12)
    run_autor.font.color.rgb = RGBColor(27, 58, 92)

    p_gen = doc.add_paragraph()
    p_gen.paragraph_format.space_after = Pt(10)

    run_gen = p_gen.add_run("GENERO / FORMATO:    _________________________________________________")
    run_gen.bold = True
    run_gen.font.size = Pt(12)
    run_gen.font.color.rgb = RGBColor(27, 58, 92)

    # === SECCION A ===
    add_section_header(doc, "A", "El Origen del Manuscrito")
    add_horizontal_line(doc)

    add_question(
        doc, "A.1",
        "¿Que te impulso a escribir este libro? Describe el momento o la necesidad que dio origen al proyecto.",
        "El sistema buscara el mensaje central en el texto, pero necesita saber si coincide con la razon que te movio a escribir.",
        "Una crisis personal, una promesa a alguien, una voz interior persistente, una observacion de la sociedad..."
    )

    add_question(
        doc, "A.2",
        "Si tuvieras que explicarle a un amigo de confianza de que trata tu libro (sin frases de marketing), ¿que dirias?",
        "Captura la esencia en tus propias palabras. Esto calibra como el sistema interpretara la sinopsis automatica.",
        "Es la historia de una madre que pierde a su hijo y encuentra una razon para seguir..."
    )

    add_question(
        doc, "A.3",
        "¿Hay una experiencia personal, investigacion profunda o mundo vivido que alimente este texto? Describe el vinculo.",
        "La fuente del material (vida propia, entrevistas, imaginacion, estudio academico) cambia como se debe leer y editar el libro.",
        "Trabaje 10 anos como maestra rural y vi cosas que nadie escribe... / Es pura invencion, pero basada en la mitologia nordica..."
    )

    add_question(
        doc, "A.4",
        "¿Que parte del manuscrito te costo mas escribir? ¿Y cual fue la que fluyo con naturalidad?",
        "Identifica zonas de resistencia (que quizas necesitan mas claridad) y zonas de fuerza (que pueden expandirse o servir de modelo).",
        "El capitulo del duelo me tomo seis meses... / Los dialogos entre hermanos salieron solos en una noche..."
    )

    add_question(
        doc, "A.5",
        "¿Hay algo en el texto que hoy, con perspectiva, ya no representa del todo lo que sientes o piensas?",
        "Detecta material que podria necesitar actualizacion, advertencia al lector, o reescritura antes de la edicion final.",
        "El personaje secundario Maria ahora me parece un estereotipo... / El final que escribi enojado ya no me representa..."
    )

    add_separator(doc)

    # === SECCION B ===
    add_section_header(doc, "B", "El Lector que Tu Imaginaste")
    add_horizontal_line(doc)

    add_question(
        doc, "B.1",
        "Describe a la persona que tenias en mente mientras escribias. ¿Quien es? ¿Que esta viviendo? ¿Que necesita?",
        "El sistema generara un perfil de lector automaticamente, pero ese perfil puede ser generico. Tu vision del lector real es oro.",
        "Un joven de 25 anos que dejo la iglesia y busca respuestas sin condena... / Una abuela que quiere dejar su historia escrita..."
    )

    add_question(
        doc, "B.2",
        "¿Has recibido feedback de lectores beta, editores, o personas de confianza? ¿Que te dijeron que recordaste?",
        "Los comentarios reales revelan si el mensaje esta llegando, si hay confusion, o si hay pasajes que conectan especialmente.",
        "Mi hermana dijo que el capitulo 3 le hizo sentir vista por primera vez... / Tres personas me preguntaron que paso con el personaje X..."
    )

    add_question(
        doc, "B.3",
        "¿Que deberia sentir, pensar o hacer una persona justo despues de terminar el ultimo capitulo?",
        "El analisis detectara emociones en el texto, pero tu sabes que transformacion querias provocar intencionalmente.",
        "Deberia sentir que no esta sola... / Deberia cerrar el libro y llamar a su padre... / Deberia querer releerlo..."
    )

    add_question(
        doc, "B.4",
        "¿Hay un tipo de lector para quien este libro NO esta pensado? ¿Por que?",
        "Saber a quien no va dirigido es tan util como saber a quien si. Evita malentendidos y refina el tono.",
        "No es para quien busca un manual paso a paso... / No es para lectores que odian saltos temporales..."
    )

    add_separator(doc)

    # === SECCION C ===
    add_section_header(doc, "C", "La Arquitectura del Texto")
    add_horizontal_line(doc)

    add_question(
        doc, "C.1",
        "¿Como describirias la estructura de tu libro? (lineal, circular, con multiples voces, cronologico, fragmentado, etc.)",
        "El sistema mapeara capitulos automaticamente, pero necesita saber si la estructura que ve es la que tu diseñaste.",
        "Tres generaciones contadas hacia atras... / Alterno entre el presente y diarios de 1980... / Es un solo monologo interior..."
    )

    add_question(
        doc, "C.2",
        "¿Hay momentos clave en la narrativa que consideras puntos de no retorno, revelaciones o cambios de direccion?",
        "Ayuda al sistema a no confundir giros intencionales con problemas estructurales.",
        "En la pagina 45 el protagonista descubre la carta... / El capitulo 8 cambia todo lo que el lector creia saber..."
    )

    add_question(
        doc, "C.3",
        "¿Hay capítulos o secciones que dudes si deberían quedarse, irse, o fusionarse?",
        "El sistema detectara ritmos irregulares, pero no sabra cuales secciones tu ya sospechas que sobran o faltan.",
        "El prologo es largo y no se si necesario... / Los dos ultimos capitulos se sienten apresurados..."
    )

    add_question(
        doc, "C.4",
        "¿Hay un capitulo, escena o idea que te gustaria agregar y aun no has escrito?",
        "El sistema podra recomendar expansiones, pero necesita saber si el manuscrito esta cerrado o aun respira.",
        "Falta una escena donde el padre explique por que se fue... / Quiero agregar un mapa al principio..."
    )

    add_question(
        doc, "C.5",
        "¿Cual es la longitud actual del manuscrito (paginas o palabras)? ¿Consideras que esta completo o le falta cuerpo?",
        "La extension influye en el tipo de edicion recomendada. Un libro corto a proposito se edita diferente a uno que necesita expansion.",
        "Tiene 120 paginas y es un ensayo corto a proposito... / Son 200 paginas pero siento que la trama necesita mas desarrollo..."
    )

    add_separator(doc)

    # === SECCION D ===
    add_section_header(doc, "D", "La Voz y el Estilo")
    add_horizontal_line(doc)

    add_question(
        doc, "D.1",
        "¿Como describirias tu voz como autor en este libro? (intima, distante, ironica, poetica, didactica, coloquial, etc.)",
        "El analizara patrones linguisticos, pero tu percepcion consciente de tu propia voz puede revelar intenciones ocultas.",
        "Hablo como si le contara a mi hermano menor... / Uso un lenguaje formal porque el tema lo amerita..."
    )

    add_question(
        doc, "D.2",
        "¿Hay recursos o patrones que uses a proposito y que un editor inexperto podria confundir con errores?",
        "Frases cortas, repeticiones, falta de signos de puntuacion, mezcla de tiempos verbales, dialecto — todo es valido si es intencional.",
        "Repito la frase \"y entonces\" como mantra narrativo a proposito... / Dejo oraciones sin terminar para crear tension..."
    )

    add_question(
        doc, "D.3",
        "¿Con que autor o estilo te gustaria que tu libro dialogara o conviviera en una estanteria?",
        "El sistema generara comparables, pero tus referentes conscientes calibran mejor el universo literario al que perteneces.",
        "Me gustaria que estuviera cerca de Garcia Marquez en tono, pero mas corto... / Es como un Bukowski sin el autodestruccion..."
    )

    add_separator(doc)

    # === SECCION E ===
    add_section_header(doc, "E", "Lo que el Sistema Debe Saber")
    add_horizontal_line(doc)

    add_question(
        doc, "E.1",
        "¿Hay algo en tu formacion, oficio, o trayectoria personal que el sistema deberia tener en cuenta al leerte?",
        "Un medico escribe diferente a un poeta. Un pastor escribe diferente a un novelista. Tu contexto personal enriquece la lectura.",
        "Soy veterinario, no humanista, asi que mis metaforas vienen del mundo animal... / Soy abogado y eso hace que estructure todo en argumentos..."
    )

    add_question(
        doc, "E.2",
        "¿Hay temas delicados (muerte, violencia, religion, politica, identidad) que requieren sensibilidad especial en el analisis?",
        "El sistema detectara palabras clave, pero no el peso emocional o cultural que tu sabes que tienen para tu lector.",
        "El libro habla del suicidio de mi hermano... hay pasajes que son dolorosos y no quiero que los marquen como \"demasiado melodramaticos\"..."
    )

    add_question(
        doc, "E.3",
        "¿Que prioridad tiene para ti: la belleza del lenguaje, la claridad del mensaje, o la fuerza emocional de la narrativa?",
        "Esta respuesta le dice al sistema que tipo de recomendaciones debe privilegiar: estilisticas, estructurales, o de impacto.",
        "Prefiero que sea hermoso aunque no se entienda del todo... / Lo importante es que el lector no se pierda nunca..."
    )

    add_question(
        doc, "E.4",
        "¿Hay algo que NO te pregunte y que crees que el sistema deberia saber antes de analizar tu manuscrito?",
        "Un espacio abierto para lo impredecible. La mejor calibracion viene de lo que no se nos ocurrio preguntar.",
        "El libro tiene dos finales posibles y aun no decido cual usar... / Esta basado en hechos reales pero cambie los nombres..."
    )

    # === CIERRE ===
    doc.add_paragraph()
    p_cierre = doc.add_paragraph()
    p_cierre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cierre.paragraph_format.space_before = Pt(30)
    p_cierre.paragraph_format.space_after = Pt(10)

    run_cierre = p_cierre.add_run("━" * 40)
    run_cierre.font.size = Pt(10)
    run_cierre.font.color.rgb = RGBColor(212, 168, 67)

    p_gracias = doc.add_paragraph()
    p_gracias.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run_gracias = p_gracias.add_run("Gracias por confiar tu manuscrito a este proceso.\n"
                                     "Tus respuestas son la brujula que guiara todo el analisis.")
    run_gracias.font.size = Pt(12)
    run_gracias.font.color.rgb = RGBColor(27, 58, 92)
    run_gracias.font.name = "Calibri"

    p_instr = doc.add_paragraph()
    p_instr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_instr.paragraph_format.space_before = Pt(10)

    run_instr = p_instr.add_run("Guarda tus respuestas junto al manuscrito antes de ejecutar el pipeline.")
    run_instr.font.size = Pt(10)
    run_instr.font.color.rgb = RGBColor(100, 100, 100)
    run_instr.font.name = "Calibri"

    output_path = r"c:\Users\crist\OneDrive\Desktop\Claude\TSBN\Catalogo\CUESTIONARIO_GENERICO_ESCRITORES.docx"
    doc.save(output_path)
    print(f"✅  Cuestionario generico generado: {output_path}")

if __name__ == "__main__":
    main()
