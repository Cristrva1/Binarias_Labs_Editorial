"""
Generador de Cuestionario TSBN en Word — Diseño profesional para Arturo Ledezma
Paletas: Dorado/Ocre + Azul Marino (elegante, legible, cálido)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color_hex):
    """Aplica color de fondo a una celda de tabla."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def add_horizontal_line(doc, color_hex="D4A843", thickness=8):
    """Añade una línea decorativa."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("\n" + "━" * 60 + "\n")
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_section_header(doc, number, title, color_hex="1B3A5C"):
    """Añade un encabezado de sección con barra de color."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.page_break_before = False

    run = p.add_run(f"SECCION {number}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16))
    run.font.name = "Calibri"

    run2 = p.add_run(f"  —  {title.upper()}")
    run2.bold = True
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16))
    run2.font.name = "Calibri"

def add_question(doc, number, question, contexto, ejemplo=""):
    """Añade una pregunta con contexto y espacio para respuesta."""
    # Número y pregunta
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)

    run_num = p.add_run(f"{number}.  ")
    run_num.bold = True
    run_num.font.size = Pt(12)
    run_num.font.color.rgb = RGBColor(212, 168, 67)  # Dorado
    run_num.font.name = "Calibri"

    run_q = p.add_run(question)
    run_q.bold = True
    run_q.font.size = Pt(12)
    run_q.font.color.rgb = RGBColor(27, 58, 92)  # Azul marino
    run_q.font.name = "Calibri"

    # Contexto (cursiva, gris)
    p_ctx = doc.add_paragraph()
    p_ctx.paragraph_format.space_before = Pt(2)
    p_ctx.paragraph_format.space_after = Pt(6)
    p_ctx.paragraph_format.left_indent = Inches(0.25)

    run_ctx = p_ctx.add_run(f"💡  {contexto}")
    run_ctx.italic = True
    run_ctx.font.size = Pt(10)
    run_ctx.font.color.rgb = RGBColor(100, 100, 100)
    run_ctx.font.name = "Calibri"

    # Ejemplo (si aplica)
    if ejemplo:
        p_ej = doc.add_paragraph()
        p_ej.paragraph_format.space_before = Pt(2)
        p_ej.paragraph_format.space_after = Pt(4)
        p_ej.paragraph_format.left_indent = Inches(0.25)

        run_ej = p_ej.add_run(f"Ejemplo: {ejemplo}")
        run_ej.italic = True
        run_ej.font.size = Pt(10)
        run_ej.font.color.rgb = RGBColor(120, 100, 60)
        run_ej.font.name = "Calibri"

    # Líneas guía para respuesta
    p_resp = doc.add_paragraph()
    p_resp.paragraph_format.space_before = Pt(6)
    p_resp.paragraph_format.space_after = Pt(4)

    run_label = p_resp.add_run("Tu respuesta:")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_label.font.color.rgb = RGBColor(150, 150, 150)
    run_label.font.name = "Calibri"

    # Líneas para escribir
    for _ in range(3):
        p_line = doc.add_paragraph()
        p_line.paragraph_format.space_before = Pt(2)
        p_line.paragraph_format.space_after = Pt(2)
        p_line.paragraph_format.left_indent = Inches(0.25)

        run_line = p_line.add_run("_" * 90)
        run_line.font.size = Pt(10)
        run_line.font.color.rgb = RGBColor(200, 200, 200)
        run_line.font.name = "Calibri"

def add_separator(doc):
    """Añade separador visual."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run("✦  ✦  ✦")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(212, 168, 67)  # Dorado
    run.font.name = "Calibri"

def main():
    doc = Document()

    # Configurar márgenes amplios para comodidad
    sections = doc.sections[0]
    sections.top_margin = Cm(2.5)
    sections.bottom_margin = Cm(2.5)
    sections.left_margin = Cm(2.5)
    sections.right_margin = Cm(2.5)

    # === PORTADA ===
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)

    run = p.add_run("CUESTIONARIO PARA EL AUTOR")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(27, 58, 92)  # Azul marino
    run.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(12)

    run2 = p2.add_run("Todas Son Buenas Noticias")
    run2.italic = True
    run2.font.size = Pt(18)
    run2.font.color.rgb = RGBColor(212, 168, 67)  # Dorado
    run2.font.name = "Calibri"

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(30)

    run3 = p3.add_run("Arturo Ledezma Ruan")
    run3.font.size = Pt(14)
    run3.font.color.rgb = RGBColor(80, 80, 80)
    run3.font.name = "Calibri"

    # Tabla informativa
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F0E6")  # Crema claro

    p_info = cell.paragraphs[0]
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_info.paragraph_format.space_before = Pt(8)
    p_info.paragraph_format.space_after = Pt(8)

    run_info = p_info.add_run(
        "📖  Objetivo: Este formulario se completa ANTES de que los equipos de análisis trabajen.\n"
        "Tus respuestas calibrarán todo el pipeline editorial, de marketing y de refinamiento.\n\n"
        "⏱  Tiempo estimado: 20-30 minutos\n"
        "✍  Responde con la extensión que necesites. Más detalle = mejor resultado.\n"
        "💡 No hay respuestas incorrectas. Si una pregunta no aplica, escribe \"N/A\"."
    )
    run_info.font.size = Pt(11)
    run_info.font.color.rgb = RGBColor(60, 60, 60)
    run_info.font.name = "Calibri"

    # Borde de tabla
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()  # Espacio

    # === SECCION A ===
    add_section_header(doc, "A", "El Libro: Origen e Intención")
    add_horizontal_line(doc)

    add_question(
        doc, "A.1",
        "¿Por qué escribiste este libro? ¿Qué te impulsó a sentarte y empezar?",
        "El sistema necesita saber si el mensaje que detecte coincide con tu intención real. También busca capas emocionales o espirituales que un análisis automático puede pasar por alto.",
        "Quería dejar algo para mis hijos... / Dios me habló en una temporada difícil... / La gente me decía que tenía que escribirlo..."
    )

    add_question(
        doc, "A.2",
        "Si tuvieras que explicarle a alguien en un elevador de qué trata el libro (sin usar frases de marketing), ¿qué dirías?",
        "Ayuda al sistema a capturar la esencia del mensaje en tus propias palabras, no en lenguaje publicitario.",
        "Es sobre cómo encontrar paz cuando todo parece estar en contra..."
    )

    add_question(
        doc, "A.3",
        "¿Existe una experiencia personal específica (crisis, milagro, revelación) que dio origen al libro? ¿Cuál?",
        "El sistema detectará temas, pero no sabrá si vienen de una experiencia vivida, de un estudio bíblico profundo, o de tu trabajo con personas. Eso cambia el ángulo de marketing y la voz del autor.",
        "Perdí mi trabajo en 2020 y en esa crisis descubrí..."
    )

    add_question(
        doc, "A.4",
        "¿Qué parte del libro te costó más escribir? ¿Y cuál fue la más disfrutable?",
        "Ayuda a identificar resistencias del autor y fortalezas genuinas que se pueden potenciar.",
        "El capítulo 5 fue difícil porque hablaba de mi padre... / Me encantó escribir las historias del capítulo 3..."
    )

    add_question(
        doc, "A.5",
        "¿Hay algo en el libro que hoy, con perspectiva, ya no representa del todo lo que crees o sientes?",
        "Detectar material que podría necesitar actualización o reescritura en próximas ediciones.",
        "El capítulo 2 lo escribí hace 3 años y ahora pienso diferente sobre..."
    )

    add_separator(doc)

    # === SECCION B ===
    add_section_header(doc, "B", "El Lector: ¿A Quién Imaginaste?")
    add_horizontal_line(doc)

    add_question(
        doc, "B.1",
        "Describe a la persona que tenías en mente mientras escribías. ¿Quién era? ¿Qué estaba viviendo?",
        "El sistema generará un buyer persona automáticamente, pero ese perfil puede ser genérico. Necesitamos al lector REAL que tú imaginaste.",
        "Mi sobrina de 28 años que acaba de divorciarse y busca sentido... / Un hombre de 50 que perdió su negocio..."
    )

    add_question(
        doc, "B.2",
        "¿Has recibido feedback de lectores reales? ¿Qué te dijeron?",
        "Testimonios reales son oro para marketing. También revelan si el mensaje está llegando o no.",
        "Mi pastor dijo que el capítulo 4 le hizo llorar... / Una amiga me dijo que no entendió el capítulo 7..."
    )

    add_question(
        doc, "B.3",
        "¿Hay un tipo de lector que NO quieres que lea este libro? ¿Por qué?",
        "Tanto como saber a quién va dirigido, es útil saber a quién NO va dirigido. Refina el posicionamiento.",
        "No es para quien busca un manual técnico de teología..."
    )

    add_question(
        doc, "B.4",
        "¿Qué debería sentir o pensar una persona justo después de terminar el último capítulo?",
        "El análisis detectará emociones dominantes por capítulo, pero tú como autor sabes qué transformación querías provocar.",
        "Debería sentir esperanza y tener ganas de orar... / Debería entender que Dios tiene un plan..."
    )

    add_question(
        doc, "B.5",
        "¿Esperas que el lector haga algo concreto después de leer el libro (acción, cambio, oración, contactarte)?",
        "Si hay un llamado a la acción implícito, el sistema puede potenciarlo en el marketing.",
        "Quiero que me escriba a mi email... / Quiero que busque una iglesia... / Quiero que comparta el libro..."
    )

    add_separator(doc)

    # === SECCION E ===
    add_section_header(doc, "E", "Dudas y Decisiones Pendientes")
    add_horizontal_line(doc)

    add_question(
        doc, "E.1",
        "¿Hay capítulos o secciones que dudes si deberían quedarse, irse, o fusionarse?",
        "El sistema detectará problemas estructurales, pero no sabrá cuáles secciones son \"sagradas\" para ti vs cuáles estás dispuesto a cambiar.",
        "El capítulo 6 es muy corto, no sé si fusionarlo con el 5..."
    )

    add_question(
        doc, "E.2",
        "¿Hay un capítulo o idea que te gustaría agregar y aún no has escrito?",
        "El pipeline podrá recomendar expansiones, pero necesita saber si el libro está definitivamente cerrado o aún en evolución.",
        "Me falta escribir sobre cómo perdonar a quienes te traicionan..."
    )

    add_question(
        doc, "E.3",
        "El libro tiene 91 páginas. ¿Consideras que está completo, o es una versión abreviada de algo más largo?",
        "91 páginas es delgado para autoayuda espiritual. ¿Es un libro corto a propósito, o falta material que aún planeas agregar?",
        "Es un libro corto a propósito, como una carta larga... / Sí quiero agregar 3 capítulos más..."
    )

    add_question(
        doc, "E.4",
        "¿Te preocupa más la calidad literaria (prosa, estilo), la profundidad teológica, o el potencial de ventas?",
        "Esta respuesta define qué equipo del pipeline tendrá más peso en sus recomendaciones.",
        "Lo más importante es que la gente sea tocada... / Quiero que sea un bestseller... / Quiero que sea bíblicamente sólido..."
    )

    add_question(
        doc, "E.5",
        "¿Hay algo que el sistema debería SABER sobre ti como autor para no malinterpretar tu estilo?",
        "El análisis automático podría marcar como \"debilidad\" algo que es intencional. Por ejemplo: frases cortas como estilo poético, repeticiones como recurso retórico, etc.",
        "Uso frases cortas a propósito porque vengo de la poesía... / Repito palabras como recurso de énfasis..."
    )

    add_separator(doc)

    # === SECCION F ===
    add_section_header(doc, "F", "Espacio Libre")
    add_horizontal_line(doc)

    add_question(
        doc, "F.1",
        "¿Qué no te pregunté y debería haber preguntado?",
        "Esta pregunta mejora las próximas versiones del cuestionario y del pipeline.",
        "Nadie me preguntó sobre mi relación con mi coautor... / Deberían haber preguntado si hay otro libro en camino..."
    )

    # === CIERRE ===
    doc.add_paragraph()
    p_cierre = doc.add_paragraph()
    p_cierre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cierre.paragraph_format.space_before = Pt(30)
    p_cierre.paragraph_format.space_after = Pt(10)

    run_cierre = p_cierre.add_run("━" * 40)
    run_cierre.font.size = Pt(10)
    run_cierre.font.color.rgb = RGBColor(212, 168, 67)

    p_gracias = doc.add_paragraph()
    p_gracias.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run_gracias = p_gracias.add_run("¡Gracias por tu tiempo, Arturo!\nTus respuestas son el corazón de este proyecto.")
    run_gracias.font.size = Pt(12)
    run_gracias.font.color.rgb = RGBColor(27, 58, 92)
    run_gracias.font.name = "Calibri"

    p_instr = doc.add_paragraph()
    p_instr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_instr.paragraph_format.space_before = Pt(10)

    run_instr = p_instr.add_run("Guarda tus respuestas en:  RESPUESTAS_AUTOR_TSBN.md")
    run_instr.font.size = Pt(10)
    run_instr.font.color.rgb = RGBColor(100, 100, 100)
    run_instr.font.name = "Calibri"

    # Guardar
    output_path = r"c:\Users\crist\OneDrive\Desktop\Claude\TSBN\Catalogo\CUESTIONARIO_TSBN_ARTURO.docx"
    doc.save(output_path)
    print(f"✅  Cuestionario generado: {output_path}")

if __name__ == "__main__":
    main()
